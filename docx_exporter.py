"""
DOCX Exporter Module
Exports resume data to Microsoft Word format.
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import base64
from utils.font_mapper import get_font_substitute

class DOCXExporter:
    def __init__(self, layout_data):
        """
        Initialize DOCX exporter with layout information.
        
        Args:
            layout_data (dict): Layout information from parsed PDF
        """
        self.layout = layout_data
        self.doc = Document()
        
        # Set page margins based on PDF layout
        margins = layout_data.get("margins", {})
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(margins.get("top", 50) / 72)
            section.bottom_margin = Inches(margins.get("bottom", 50) / 72)
            section.left_margin = Inches(margins.get("left", 50) / 72)
            section.right_margin = Inches(margins.get("right", 50) / 72)
    
    def export(self, fields, images):
        """
        Export resume data to DOCX format.
        
        Args:
            fields (list): List of text fields with content and formatting
            images (list): List of images with position and data
            
        Returns:
            bytes: DOCX file as bytes
        """
        # Sort fields by vertical position (top to bottom)
        sorted_fields = sorted(fields, key=lambda f: f["position"]["y"])
        
        # Group fields by approximate Y position (same line)
        lines = self._group_fields_by_line(sorted_fields)
        
        # Add images at the top if they exist
        for img in images:
            if img["position"]["y"] < 200:  # Top section images (like profile photo)
                self._add_image(img)
        
        # Add text content
        for line_fields in lines:
            self._add_line(line_fields)
        
        # Add remaining images
        for img in images:
            if img["position"]["y"] >= 200:
                self._add_image(img)
        
        # Save to bytes
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        
        return buffer.getvalue()
    
    def _group_fields_by_line(self, fields, threshold=5):
        """
        Group fields that appear on the same line.
        
        Args:
            fields (list): Sorted list of fields
            threshold (int): Y-position threshold for grouping
            
        Returns:
            list: List of field groups (lines)
        """
        if not fields:
            return []
        
        lines = []
        current_line = [fields[0]]
        current_y = fields[0]["position"]["y"]
        
        for field in fields[1:]:
            y = field["position"]["y"]
            if abs(y - current_y) <= threshold:
                # Same line
                current_line.append(field)
            else:
                # New line
                lines.append(current_line)
                current_line = [field]
                current_y = y
        
        # Add last line
        if current_line:
            lines.append(current_line)
        
        return lines
    
    def _add_line(self, fields):
        """
        Add a line of text fields to the document.
        
        Args:
            fields (list): Fields on the same line
        """
        # Sort fields by X position (left to right)
        sorted_fields = sorted(fields, key=lambda f: f["position"]["x"])
        
        # Create paragraph
        paragraph = self.doc.add_paragraph()
        
        for field in sorted_fields:
            # Add run with formatting
            run = paragraph.add_run(field["value"])
            
            # Apply font formatting
            font_info = field["font"]
            font_name, _ = get_font_substitute(font_info.get("family", "Arial"))
            
            run.font.name = font_name
            run.font.size = Pt(font_info.get("size", 12))
            run.font.bold = font_info.get("bold", False)
            run.font.italic = font_info.get("italic", False)
            
            # Apply color
            color = font_info.get("color", "#000000")
            rgb = self._hex_to_rgb(color)
            run.font.color.rgb = RGBColor(rgb[0], rgb[1], rgb[2])
            
            # Add space between fields on same line
            if field != sorted_fields[-1]:
                paragraph.add_run("  ")
    
    def _add_image(self, image):
        """
        Add an image to the document.
        
        Args:
            image (dict): Image data with position and base64 content
        """
        # Decode base64 image
        image_data = base64.b64decode(image["data"])
        image_buffer = io.BytesIO(image_data)
        
        # Add image to document
        # Convert PDF points to inches
        width_inches = image["position"]["width"] / 72
        height_inches = image["position"]["height"] / 72
        
        # Add paragraph for image
        paragraph = self.doc.add_paragraph()
        run = paragraph.add_run()
        
        try:
            run.add_picture(image_buffer, width=Inches(width_inches))
        except Exception as e:
            # If image can't be added, skip it
            print(f"Warning: Could not add image - {str(e)}")
    
    def _hex_to_rgb(self, hex_color):
        """
        Convert hex color to RGB tuple (0-255 range).
        
        Args:
            hex_color (str): Hex color code (e.g., "#FF0000")
            
        Returns:
            tuple: RGB values in 0-255 range
        """
        hex_color = hex_color.lstrip('#')
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return (r, g, b)
